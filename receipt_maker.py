import csv,operator
from datetime import date
from docx import Document
from docx2pdf import convert;
import os;
from docxcompose.composer import Composer
import shutil;


def getMonthText():
    today = date.today();
    months = ['January','February','March','April','May','June','July','August','September','October','November','December']
    return months[today.month-1]


def find_csv_filenames():
    filenames = os.listdir('CSVs/')
    return [ filename for filename in filenames if filename.endswith('.csv') ]

def GetInvoiceID(num):
    today = date.today();
    tMonth = str(today.month)
    tNum = str(num);
    if (len(tMonth)==1):
        tMonth = '0'+tMonth
    if (len(tNum)==1):
        tNum = '000' + tNum
    if (len(tNum)==2):
        tNum = '00' + tNum
    if (len(tNum)==3):
        tNum = '0' + tNum
    today = ('{}-{}'.format(today.year,tMonth,));
    return ('INV-{}-{}'.format(today,tNum));


def GetCommissionData(file):
    print(file)
    file = open(file,'r');
    comData = csv.reader(file,delimiter=',')
    comData = sorted(comData, key=operator.itemgetter(2))
    file.close();
    #GET ALL INFO FROM THE BRANDS AND TURN IT INTO A JSON FILE
    brands = [];
    for d in comData:
        brands.append(d[2]);
    brands = list(dict.fromkeys(brands)) #GET RID OF DUPLICATES IN LIST
    brands.remove('BrandID') #GETS RID OF BRAND TITLE IN LIST FILE
    allInvoices = [];  #LIST OF ALL INVOICES
    for b in brands:
        amount = 0;
        orders = 0;
        id = GetInvoiceID((brands.index(b)) + 1);
        invoice = [];
        for d in comData:
            if (d[2] == b):
                amount+=float(d[3]);
                orders+=1;
                invoice.append({
                    "Amount": round(float(d[3]),2),
                    "Date": d[4],
                    "Description": d[1],
                })#ADDING BRAND VERSION TO COMPLETE INVOICE JSON
        allInvoices.append({
            'invoiceID': id,
            "BrandID": b,
            "Amount": round(amount,2),
            "Number": orders,
            "Orders": invoice
        })
    return allInvoices;

def GetBusinessData(file):
    file = open(file,'r');
    comData = csv.reader(file,delimiter=',')
    comData = sorted(comData, key=operator.itemgetter(2))
    file.close();
    #GET ALL INFO FROM THE BRANDS AND TURN IT INTO A JSON FILE
    bus = [];
    for d in comData:
        try:
            add = d[4].split(',')
            try:
                add.remove('Address');
            except: 
                None
            js = {
                'CustomerID':d[0],
                'BrandID': d[1],
                'Name': d[2],
                'Email':d[3],
                'Address': {
                    'Street': add[0],
                    'Town': add[1],
                    'County': add[2],
                    'Postcode': add[3],
                }
            }
            bus.append(js)
        except:
            None
    return bus;

def GetDiscountData(file):
    file = open(file,'r');
    comData = csv.reader(file,delimiter=',')
    comData = sorted(comData, key=operator.itemgetter(2))
    file.close();
    #GET ALL INFO FROM THE BRANDS AND TURN IT INTO A JSON FILE
    discounts = [];
    for d in comData:
        try:
            js = {
                'discountID':d[1],
                'description':d[3],
            }
            discounts.append(js)
        except:
            None
    return discounts;

def getBusiness(id,list):
    for l in list:
        if (l['BrandID']==id):
            return l;


def getDescription(id,dis):
    for l in dis:
        if (id==l['discountID']):
            return l['description'];


def editWordDocument(invoice,info):
    #replaces text based on what is set in both the variables and Template document.
    template = Document('Templates/Receipt template.docx')
    output = ('Temp/Top-%s.docx'%(str(invoice['invoiceID'])))
    dates = date.today();
    next = int(dates.month) + 1;
    tNum = str(dates.month)
    nNum = str(next);
    tYear = str(dates.year).replace('20','');
    if (len(tNum)==1):
        tNum = '0'+tNum
    if (len(nNum) == 1):
        nNum = '0'+nNum
    today = ('{}-{}-{}'.format(dates.day,tNum,tYear));
    nMonth = ('{}-{}-{}'.format(dates.day,nNum,tYear));
    tNum = str(dates.month);
    variables = {
        "COMPANY_NAME": info['Name'],
        "ADDRESS_STREET": info['Address']['Street'],
        "ADDRESS_TOWN": info['Address']['Town'],
        "ADDRESS_COUNTY": info['Address']['County'],
        "ADDRESS_POSTCODE": info['Address']['Postcode'],
        'DATEISSUED': today,
        'DATEDUE': nMonth,
        'TOTALPAYMENT': ('£' + str(invoice['Amount'])),
        'IDNUMBER': invoice['invoiceID']
    }
    for variable_key, variable_value in variables.items():
        for paragraph in template.paragraphs:
            replacement_text(paragraph, variable_key,variable_value);
        for table in template.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replacement_text(paragraph,variable_key,variable_value);
    template.save(output);

def replacement_text(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


def addWordTable(invoice,discounts):
    #THIS creates a table in a word document
    doc = Document();
    data = [];
    for o in invoice['Orders']:
        entry = (o['Description'],o['Date'],o['Amount'],)
        data.append(entry);
    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0].cells
    row[0].text = 'Description'
    row[1].text = 'Transaction date'
    row[2].text = 'Commission'
    for des,date,comm in data:
        row = table.add_row().cells
        desc = getDescription(des,discounts);
        row[0].text = desc,
        row[1].text = date,
        row[2].text = ('£'+str(comm),)
    #table.style = 'Medium Shading 1 Accent 6'
    row[0].text = ' '
    row[1].text = 'Total'
    row[2].text = ('£' + str(invoice['Amount']))
    path = 'Temp/Table-%s.docx'%invoice['invoiceID'];
    
    doc.save(path);

def mergeWordDocument(files,inv):
    result = Document(files[0])
    composer = Composer(result)
    
    doc1 = Document(files[1]);
    composer.append(doc1);
    changed = 'Temp/%s.docx'%inv['invoiceID']
    composer.save(changed)
    

def makeDocPDF(word,pdf):
    convert(word,pdf)









# WHERE ALL THE STUFF IS RUN
files = find_csv_filenames();
allInvoices = GetCommissionData('CSVs/%s'%files[0])
allBusinesses = GetBusinessData('CSVs/%s'%files[2])
allDiscounts = GetDiscountData('CSVs/%s'%files[1])
today = date.today();
year = today.year
month = getMonthText();
for inv in allInvoices:
    business = getBusiness(inv['BrandID'],allBusinesses)
    try:
        os.mkdir('Temp');
    except:
        None
    try:
        os.mkdir('Receipt');
    except:
        None
    try:
        os.mkdir('Receipt/%s'%year);
    except:
        None
    try:
        os.mkdir('Receipt/%s/%s'%(year,month));
    except:
        None
    editWordDocument(inv,business);
    addWordTable(inv,allDiscounts);
    top = ('Temp/Top-%s.docx'%(str(inv['invoiceID'])))
    table = ('Temp/Table-%s.docx'%(str(inv['invoiceID'])))
    files = [top,table,'Templates/Invoice_PayTab.docx']
    mergeWordDocument(files,inv)
    word = 'Temp/%s.docx'%inv['invoiceID']
    pdf = ('Receipt/%s/%s/%s.pdf'%(year,month,inv['invoiceID']))
    makeDocPDF(word,pdf)
shutil.rmtree('Temp');

print('DONE');
