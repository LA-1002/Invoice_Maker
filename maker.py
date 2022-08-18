import csv,operator
from datetime import date
from docx import Document
from docx2pdf import convert;
import os;
from docxcompose.composer import Composer
import shutil



def find_csv_filenames():
    filenames = os.listdir('CSVs/')
    return [ filename for filename in filenames if filename.endswith('.csv') ]

def GetInvoiceID(num):
    today = date.today();
    tMonth = str(today.month)
    tNum = str(num);
    if (len(tMonth)==1):
        tMonth = '0'+tMonth
    if (len(tNum)==1):
        tNum = '000' + tNum
    if (len(tNum)==2):
        tNum = '00' + tNum
    if (len(tNum)==3):
        tNum = '0' + tNum
    today = ('{}-{}'.format(today.year,tMonth,));
    return ('INV-{}-{}'.format(today,tNum));
def GetCommissionData(file):
    print(file)
    file = open(file,'r');
    comData = csv.reader(file,delimiter=',')
    comData = sorted(comData, key=operator.itemgetter(2))
    file.close();
    #GET ALL INFO FROM THE BRANDS AND TURN IT INTO A JSON FILE
    brands = [];
    for d in comData:
        brands.append(d[2]);
    brands = list(dict.fromkeys(brands)) #GET RID OF DUPLICATES IN LIST
    brands.remove('BrandID') #GETS RID OF BRAND TITLE IN LIST FILE
    allInvoices = [];  #LIST OF ALL INVOICES
    for b in brands:
        amount = 0;
        orders = 0;
        id = GetInvoiceID((brands.index(b)) + 1);
        invoice = [];
        for d in comData:
            if (d[2] == b):
                amount+=float(d[3]);
                orders+=1;
                invoice.append({
                    "Amount": round(float(d[3]),2),
                    "Date": d[4],
                    "Description": "Testing"
                })#ADDING BRAND VERSION TO COMPLETE INVOICE JSON
        allInvoices.append({
            'invoiceID': id,
            "BrandID": b,
            "Amount": round(amount,2),
            "Number": orders,
            "Orders": invoice
        })
    return allInvoices;


def editWordDocument(invoice,info):
    #replaces text based on what is set in both the variables and Template document.
    template = Document('Templates/Invoice Template.docx')
    output = ('Temp/Top-%s.docx'%(str(invoice['invoiceID'])))
    dates = date.today();
    next = int(dates.month) + 1;
    tNum = str(dates.month)
    nNum = str(next);
    if (len(tNum)==1):
        tNum = '0'+tNum
    if (len(nNum) == 1):
        nNum = '0'+nNum
    today = ('{}-{}-{}'.format(dates.day,tNum,dates.year));
    nMonth = ('{}-{}-{}'.format(dates.day,nNum,dates.year));
    tNum = str(dates.month);
    variables = {
        "COMPANY_NAME": info['Name'],
        "ADDRESS_STREET": info['Address']['Street'],
        "ADDRESS_TOWN": info['Address']['Town'],
        "ADDRESS_COUNTY": info['Address']['County'],
        "ADDRESS_POSTCODE": info['Address']['Postcode'],
        'DATEISSUED': today,
        'DATEDUE': nMonth,
        'TOTALPAYMENT': ('£' + str(invoice['Amount'])),
        'IDNUMBER': invoice['invoiceID']
    }
    for variable_key, variable_value in variables.items():
        for paragraph in template.paragraphs:
            replacement_text(paragraph, variable_key,variable_value);
        for table in template.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replacement_text(paragraph,variable_key,variable_value);
    template.save(output);

def replacement_text(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


def addWordTable(invoice):
    #THIS creates a table in a word document
    doc = Document();
    data = [];
    for o in invoice['Orders']:
        entry = (o['Description'],o['Date'],o['Amount'],)
        data.append(entry);

    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0].cells
    row[0].text = 'Description'
    row[1].text = 'Transaction date'
    row[2].text = 'Commission'
    for des,date,comm in data:
        row = table.add_row().cells
        row[0].text = des,
        row[1].text = date,
        row[2].text = ('£'+str(comm),)
    table.style = ("Colorful List");
    row[0].text = ' '
    row[1].text = 'Total'
    row[2].text = ('£' + str(invoice['Amount']))
    path = 'Temp/Table-%s.docx'%invoice['invoiceID'];
    doc.save(path);

def mergeWordDocument(files,inv):
    changed = 'Temp/%s.docx'%inv['invoiceID']
    result = Document(files[0])
    composer = Composer(result)
    doc1 = Document(files[1]);
    doc2 = Document(files[2]);

    composer.append(doc1);
    composer.append(doc2);

    composer.save(changed)
    

def makeDocPDF(word,pdf):
    convert(word,pdf)









# WHERE ALL THE STUFF IS RUN
files = find_csv_filenames();
file = 'CSVs/%s'%files[0]
allInvoices = GetCommissionData(file);
business = {
'Name': "Alidi's",
'Address': {
    'Street': 'Fake Street',
    'Town': 'Counterfeit Corner',
    'County': 'Phony District',
    'Postcode': 'SHAM PRE',
}
}
for inv in allInvoices:
    try:
        os.mkdir('Temp');
    except:
        None
    try:
        os.mkdir('Invoices');
    except:
        None
    editWordDocument(inv,business);
    addWordTable(inv);
    top = ('Temp/Top-%s.docx'%(str(inv['invoiceID'])))
    table = ('Temp/Table-%s.docx'%(str(inv['invoiceID'])))
    files = [top,table,'Templates/Invoice_PayTab.docx']
    mergeWordDocument(files,inv)
    word = 'Temp/%s.docx'%inv['invoiceID']
    pdf = 'Invoices/%s.pdf'%inv['invoiceID']
    makeDocPDF(word,pdf)
shutil.rmtree('Temp');
    

print('DONE');
